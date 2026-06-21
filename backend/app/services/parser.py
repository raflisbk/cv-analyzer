from io import BytesIO

try:
    import fitz

    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import langdetect

    LANGDETECT_AVAILABLE = True
except ImportError:
    LANGDETECT_AVAILABLE = False

from app.core.logging import structured_logger as logger
from app.services.ocr import EASYOCR_AVAILABLE, perform_ocr
from app.services.validation import validate_extraction_quality
from app.services.vision import VISION_AVAILABLE, extract_text_from_image, extract_text_from_scanned_pdf


class ParsingError(Exception):
    pass


class DocumentParser:
    def parse_pdf(self, content: bytes) -> tuple[str, dict]:
        if not PYMUPDF_AVAILABLE:
            msg = "PyMuPDF not installed. Install with: pip install PyMuPDF"
            raise ParsingError(msg)

        metadata = {
            "extraction_method": None,
            "quality_score": 0.0,
            "page_count": 0,
            "detected_language": None,
        }

        try:
            doc = fitz.open(stream=content, filetype="pdf")
            metadata["page_count"] = len(doc)

            text = ""
            for page in doc:
                text += page.get_text("text", sort=True)

            doc.close()

            logger.info(
                "pdf_extraction_done",
                length=len(text),
                pages=metadata["page_count"],
            )

            if len(text.strip()) >= 50:
                quality_score, _ = validate_extraction_quality(text)
                metadata["extraction_method"] = "text_extraction"
                metadata["quality_score"] = quality_score

                if quality_score >= 0.5:
                    if LANGDETECT_AVAILABLE:
                        try:
                            metadata["detected_language"] = langdetect.detect(text)
                        except Exception:
                            logger.warning("langdetect_failed")
                            metadata["detected_language"] = "unknown"
                    else:
                        metadata["detected_language"] = "unknown"

                    return text, metadata

            logger.warning(
                "extraction_insufficient",
                text_length=len(text),
            )

        except Exception as e:
            logger.warning("pdf_extraction_failed", error=str(e))

        if VISION_AVAILABLE:
            try:
                logger.info("vision_fallback_starting")
                text = extract_text_from_scanned_pdf(content)

                if len(text.strip()) >= 50:
                    quality_score, _ = validate_extraction_quality(text)
                    metadata["extraction_method"] = "vision_granite"
                    metadata["quality_score"] = quality_score

                    if quality_score >= 0.4:
                        if LANGDETECT_AVAILABLE:
                            try:
                                metadata["detected_language"] = langdetect.detect(text)
                            except Exception:
                                metadata["detected_language"] = "unknown"
                        else:
                            metadata["detected_language"] = "unknown"
                        return text, metadata

                logger.warning("vision_insufficient", text_length=len(text))
            except Exception as e:
                logger.error("vision_fallback_failed", error=str(e))

        if EASYOCR_AVAILABLE:
            try:
                logger.info("ocr_fallback_starting")
                text, confidence = perform_ocr(content)

                if len(text.strip()) >= 50:
                    quality_score, _ = validate_extraction_quality(text)
                    metadata["extraction_method"] = "ocr"
                    metadata["quality_score"] = quality_score
                    metadata["ocr_confidence"] = confidence

                    if quality_score >= 0.4:
                        if LANGDETECT_AVAILABLE:
                            try:
                                metadata["detected_language"] = langdetect.detect(text)
                            except Exception:
                                logger.warning("langdetect_failed")
                                metadata["detected_language"] = "unknown"
                        else:
                            metadata["detected_language"] = "unknown"

                        return text, metadata

                logger.error(
                    "ocr_insufficient",
                    text_length=len(text),
                    confidence=confidence,
                    exc_info=True,
                )

            except Exception as e:
                logger.error("ocr_failed", error=str(e), exc_info=True)
        else:
            logger.warning("ocr_fallback_skipped")

        msg = (
            "Failed to extract text from PDF after trying regular extraction and OCR. "
            "File may be corrupted, password-protected, or contain only images."
        )
        raise ParsingError(msg)

    def parse_docx(self, content: bytes) -> tuple[str, dict]:
        if not DOCX_AVAILABLE:
            msg = "python-docx not installed. Install with: pip install python-docx"
            raise ParsingError(msg)

        metadata = {
            "extraction_method": "docx",
            "quality_score": 0.0,
            "detected_language": None,
        }

        try:
            doc = Document(BytesIO(content))

            text = "\n".join([para.text for para in doc.paragraphs])

            logger.info(
                "docx_extraction_done",
                length=len(text),
                paragraphs=len(doc.paragraphs),
            )

            if len(text.strip()) < 50:
                msg = "DOCX extraction yielded insufficient text (< 50 characters)"
                raise ParsingError(msg)

            quality_score, _ = validate_extraction_quality(text)
            metadata["quality_score"] = quality_score

            if quality_score < 0.3:
                msg = f"Low quality DOCX extraction (score: {quality_score})"
                raise ParsingError(msg)

            if LANGDETECT_AVAILABLE:
                try:
                    metadata["detected_language"] = langdetect.detect(text)
                except Exception:
                    metadata["detected_language"] = "unknown"
            else:
                metadata["detected_language"] = "unknown"

            return text, metadata

        except ParsingError:
            raise
        except Exception as e:
            logger.error("docx_extraction_failed", error=str(e), exc_info=True)
            msg = f"Failed to extract text from DOCX: {e!s}"
            raise ParsingError(msg) from e


    def parse_image(self, content: bytes, mime_type: str = "image/jpeg") -> tuple[str, dict]:
        """Extract text from image (JPG/PNG) using Granite Vision."""
        if not VISION_AVAILABLE:
            msg = "Vision service not available — CV_ANALYZER_REPLICATE_API_TOKEN not configured or replicate not installed"
            raise ParsingError(msg)

        metadata = {
            "extraction_method": "vision_granite",
            "quality_score": 0.0,
            "detected_language": None,
        }

        try:
            text = extract_text_from_image(content, mime_type)

            if len(text.strip()) < 50:
                msg = "Vision extraction yielded insufficient text (< 50 characters)"
                raise ParsingError(msg)

            quality_score, _ = validate_extraction_quality(text)
            metadata["quality_score"] = quality_score

            if quality_score < 0.3:
                msg = f"Low quality vision extraction (score: {quality_score})"
                raise ParsingError(msg)

            if LANGDETECT_AVAILABLE:
                try:
                    metadata["detected_language"] = langdetect.detect(text)
                except Exception:
                    metadata["detected_language"] = "unknown"
            else:
                metadata["detected_language"] = "unknown"

            logger.info("image_extraction_done", chars=len(text), quality=quality_score)
            return text, metadata

        except ParsingError:
            raise
        except Exception as e:
            logger.error("image_extraction_failed", error=str(e))
            msg = f"Failed to extract text from image: {e!s}"
            raise ParsingError(msg) from e


def parse_document(content: bytes, file_type: str) -> tuple[str, dict]:
    parser = DocumentParser()
    ext = file_type.lower()

    if ext == ".pdf":
        return parser.parse_pdf(content)
    if ext in (".docx", ".doc"):
        return parser.parse_docx(content)
    if ext in (".jpg", ".jpeg"):
        return parser.parse_image(content, "image/jpeg")
    if ext == ".png":
        return parser.parse_image(content, "image/png")
    msg = f"Unsupported file type: {file_type}"
    raise ParsingError(msg)
